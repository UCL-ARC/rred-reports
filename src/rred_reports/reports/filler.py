import io
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.table import Table, _Row
from loguru import logger


class TemplateFillerException(Exception):
    """Custom exception generator for the TemplateFiller class"""

    def __init__(self, message: str):
        super().__init__(message)
        self.message = message

    def __repr__(self) -> str:
        return self.message


class TemplateFiller:
    """TemplateFiller class to modify tables in .docx files.

    Provides save/loading, table modification and input verification.
    """

    def __init__(self, template_path: Path, header_rows: list[int], table_text_style="Table Note", table_grid_style="Grid Table 1 Light"):
        """
        Create a template filler, clean up repeated columns for tables which have a single header row

        Args:
            template_path (Path): Path to the template
            header_rows (list[int]): list of the number of header rows for each table
            table_text_style (str): text style for all new rows of table
            table_grid_style (str): grid style for all tables
        """
        try:
            self.doc = Document(template_path)
        except FileNotFoundError:
            message = f"Template file not found at {template_path}"
            raise TemplateFillerException(message) from FileNotFoundError
        self.tables = self.doc.tables
        self.header_rows = header_rows
        self.clean_tables()
        self.table_text_style = self.doc.styles[table_text_style]
        self.table_grid_style = self.doc.styles[table_grid_style]

    @staticmethod
    def remove_row(table: Table, row: _Row) -> Table:
        """Remove a single row from a Table object

        Args:
            table (Table): Table object representing a table within a .docx file
            row (_Row): Row within a .docx file

        Returns:
            table (Table): Table object representing a table within a .docx file
        """

        tbl = table._tbl  # pylint: disable=protected-access
        table_row = row._tr  # pylint: disable=protected-access
        tbl.remove(table_row)
        return table

    @staticmethod
    def remove_all_rows(table: Table, header_rows=1):
        """Remove all rows in a table except for title

        Args:
            table (Table): Table within a .docx file
            header_rows (int): number rows which make up the header
        """
        total_rows = len(table.rows) - 1
        starting_row = header_rows - 1
        for row_number in range(total_rows, starting_row, -1):
            row = table.rows[row_number]
            row._element.getparent().remove(row._element)  # pylint: disable=protected-access

    @staticmethod
    def view_header(table: Table, header_rows=1) -> list[Table.row_cells]:
        """View header text
        Can't read comments with python-docx

        Args:
            table (Table): Table object representing a table within a .docx file
            header_rows (int): number rows which make up the header

        Returns:
            list[Table.row_cells]: List of table row cells
        """
        return table.row_cells(header_rows - 1)

    def clean_tables(self):
        """Perform any necessary table cleaning steps"""
        if len(self.tables) != len(self.header_rows):
            message = f"Template filler initialised with {len(self.header_rows)} header rows but document has {len(self.tables)} tables"
            raise TemplateFillerException(message)
        self._remove_duplicated_columns()

    def _remove_duplicated_columns(self):
        """Remove duplicated columns in tables.
        Observed in some test tables, some tables contain columns that appear
        to be hidden when viewed in MS Word.
        """
        updated_tables = []
        for table, header_rows in zip(self.tables, self.header_rows):
            if header_rows != 1:
                logger.trace("Working out duplicate columns with multiple header rows is very fraught, should correct template")
                updated_tables.append(table)
                continue

            table_header = self.__class__.view_header(table)
            header_text = [cell.text.strip() for cell in table_header]

            headers_df = pd.DataFrame({"duplicate_cols": header_text}).drop_duplicates(keep="first")

            headers = headers_df["duplicate_cols"].to_list()

            if headers != header_text:
                # Duplicate header found...
                self.__class__.remove_all_rows(table)
                row = table.rows[0]
                table = self.__class__.remove_row(table, row)
                table = self.doc.add_table(rows=0, cols=len(headers), style=table.style)
                table.add_row()
                for column, header in enumerate(headers):
                    table.cell(0, column).text = header

            updated_tables.append(table)
        self.tables = updated_tables

    def populate_table(self, table_index: int, data: pd.DataFrame):
        """Populate a table with the contents of a pandas dataframe

        Args:
            table_index (int): Index of the table to populate
            data (pd.DataFrame): Pandas dataframe of future table contents
        """
        self._remove_duplicated_columns()

        table = self.tables[table_index]
        header_rows = self.header_rows[table_index]
        self.remove_all_rows(table, header_rows)
        self._verify_new_rows(table, data, header_rows)

        for i in range(data.shape[0]):
            table.add_row()
            for j in range(data.shape[-1]):
                current_cell = table.cell(i + header_rows, j)
                # Ensure NA representation is shorter: for thin columns, to avoid splitting over multiple lines
                current_cell.text = str(data.values[i, j]).replace("<NA>", "NA")
                current_cell.text = str(data.values[i, j]).replace("nan", "Missing Data")
                # manually set the style of the new text, and don't break table over multiple lines
                current_paragraph = current_cell.paragraphs[0]
                current_paragraph.style = self.table_text_style
                current_paragraph.paragraph_format.keep_with_next = True

                # Center-align the table cell
                current_paragraph.alignment = WD_ALIGN_VERTICAL.CENTER

        # override the style to deal with new rows sometimes not adding borders
        table.style = self.table_grid_style

    def _verify_new_rows(self, table: Table, data: pd.DataFrame, header_rows=1):
        """Verify dimension of new rows to be added to existing table
        Pandas dataframe width should be the columnar dimension of the table
        i.e. the number of columns should match in both

        Args:
            table (Table): Table object representing a table within a .docx file
            data (pd.DataFrame): Pandas dataframe of future table contents
            header_rows int: number rows which make up the header

        Raises:
            TemplateFillerException: DataFrame dimensions do not match table
        """
        table_header = self.__class__.view_header(table, header_rows)
        contents = [cell.text.strip() for cell in table_header]
        message = (
            f"Pandas dataframe with {data.shape[-1]} columns. Table has {len(table.columns)} columns - "
            f"dimension mismatch. Table columns are {contents}."
        )
        if not data.shape[1] == len(table.columns):
            raise TemplateFillerException(message)

    def verify_tables_filled(self) -> bool:
        """Verify all cells in all table are filled

        Returns:
            True if successful

        Raises:
            TemplateFillerException if not all cells are filled
        """
        for table_idx, table in enumerate(self.tables):
            for i, _row in enumerate(table.rows):
                for j, _col in enumerate(table.columns):
                    cell_content = table.cell(i, j).text
                    if cell_content == "":
                        message = f"Empty cell ({i},{j}) (row, col) found in table {table_idx + 1}"
                        raise TemplateFillerException(message)
        return True

    def save_document(self, output_path: Path):
        """Save the document to the defined output path

        Args:
            output_path (Path): Output file path
        """
        self.doc.save(output_path)

    def report_bytes(self) -> bytes:
        """Return a bytes representation of the report

        Args:
            file_stream.getvalue() (bytes): bytes representation of report
        """

        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        return file_stream.getvalue()
